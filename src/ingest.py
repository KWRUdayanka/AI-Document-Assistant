# src/ingest.py
"""
Module: ingest.py

This script loads documents from various formats (.pdf, .docx, .txt), extracts and chunks text, generates embeddings using OpenAI, and ingests them into a Qdrant vector database collection.
"""

import os
import uuid
import fitz

from pathlib import Path
from typing import List

from dotenv import load_dotenv
from docx import Document as DocxDocument
from llama_index.core import Document
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct, Distance, VectorParams

from custom_logging import logger

load_dotenv()

DATA_DIR = Path("./../data")
COLLECTION_NAME = os.getenv("QDRANT_COLLECTION")

# Initialize Qdrant client and vector stores
try:
    qdrant_client = QdrantClient("localhost", prefer_grpc=True)
    logger.info("Successfully initialized Qdrant client and vector stores")
except Exception as e:
    logger.error(f"Failed to initialize Qdrant client or vector stores: {e}")
    raise

# Initialize embedding model and LLM
try:
    embed_model = OpenAIEmbedding(model="text-embedding-3-small")
    llm = OpenAI(model="gpt-4o-mini", temperature=0.1)
    logger.info("Successfully initialized embedding model and LLM")
except Exception as e:
    logger.error(f"Failed to initialize AI models: {e}")
    raise

# Function to load files from directory
def load_documents_from_dir():
    """
        Loads documents from the data directory.

        Returns:
            List[Document]: A list of LlamaIndex Document objects.
    """
    docs = []
    for filename in os.listdir(DATA_DIR):
        file_path = os.path.join(DATA_DIR, filename)
        ext = Path(filename).suffix.lower()

        try:
            if ext == ".pdf":
                text = extract_text_from_pdf(file_path)
            elif ext == ".docx":
                text = extract_text_from_docx(file_path)
            elif ext == ".txt":
                text = extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type: {filename}")
                continue

            docs.append(Document(
                text=text,
                metadata={
                    "filename": filename,
                    "file_path": file_path,
                    "document_id": str(uuid.uuid4())
                }
            ))
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")

    return docs

# Function to create Qdrant collection if it doesn't exist
def create_collection(collection_name: str, vector_size: int = 1536):
    """
        Creates a Qdrant collection if it doesn't exist.

        Args:
            collection_name (str): Name of the Qdrant collection.
            vector_size (int): Size of embedding vectors.
    """
    try:
        qdrant_client.get_collection(collection_name)
        logger.info(f"Collection {collection_name} already exists")
    except Exception:
        qdrant_client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE)
        )
        logger.info(f"Created collection {collection_name}")

# Function to extract text from PDF
def extract_text_from_pdf(file_path: str) -> str:
    """
        Extracts text from a PDF file.

        Args:
            file_path (str): Path to the PDF file.

        Returns:
            str: Extracted text.
    """
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file_path: str) -> str:
    """
        Extracts text from a DOCX file.

        Args:
            file_path (str): Path to the DOCX file.

        Returns:
            str: Extracted text.
    """
    doc = DocxDocument(file_path)
    paragraphs = [para.text for para in doc.paragraphs]
    return "\n".join(paragraphs)

# Function to extract text from TXT
def extract_text_from_txt(file_path: str) -> str:
    """
        Extracts text from a TXT file.

        Args:
            file_path (str): Path to the TXT file.

        Returns:
            str: Extracted text.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

# Function to chunk text into smaller segments
def chunk_text(text: str, max_tokens: int = 512) -> List[str]:
    """
        Splits a text into smaller chunks for embedding.

        Args:
            text (str): Original text.
            max_tokens (int): Max length of each chunk.

        Returns:
            List[str]: List of text chunks.
    """
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    for word in words:
        current_length += len(word) + 1
        if current_length > max_tokens:
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word) + 1
        else:
            current_chunk.append(word)
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

# Function to ingest documents
def ingest_documents(documents: List[Document], collection_name: str):
    """
        Ingests a list of documents into Qdrant after embedding.

        Args:
            documents (List[Document]): List of LlamaIndex documents.
            collection_name (str): Name of the Qdrant collection.
    """
    # Create collection
    create_collection(collection_name)

    # Process each document
    points = []
    for doc in documents:
        file_path = doc.metadata.get("file_path")
        doc_id = doc.metadata.get("document_id", str(uuid.uuid4()))
        filename = doc.metadata.get("filename")

        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            continue

        # Extract text (assuming PDF or text input)
        if file_path.endswith(".pdf"):
            text = extract_text_from_pdf(file_path)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        if not text:
            logger.warning(f"No text extracted from {file_path}")
            continue

        # Chunk text
        chunks = chunk_text(text)
        logger.info(f"Extracted {len(chunks)} chunks from {filename}")

        # Generate embeddings and prepare points
        for i, chunk in enumerate(chunks):
            try:
                embedding = embed_model.get_text_embedding(chunk)
                # Use UUID for point ID
                point_id = str(uuid.uuid4())
                payload = {
                    "text": chunk,
                    "document_id": doc_id,
                    "chunk_index": i
                }
                points.append(PointStruct(id=point_id, vector=embedding, payload=payload))
                logger.debug(f"Generated point with ID {point_id} for chunk {i} of {filename}")
            except Exception as e:
                logger.error(f"Error generating embedding for chunk {i} of {filename}: {e}")
                continue

        logger.info(f"Generated embeddings for document {doc_id}")

    # Upsert points to Qdrant
    if points:
        try:
            qdrant_client.upsert(collection_name=collection_name, points=points)
            logger.info(f"Stored {len(points)} points in {collection_name}")
        except Exception as e:
            logger.error(f"Failed to upsert points to {collection_name}: {e}")
            raise
    else:
        logger.warning(f"No points to store for {collection_name}")

def build_index():
    """
        Loads documents from disk and ingests them into Qdrant.
    """
    logger.info(f"Reading documents...")
    doc = load_documents_from_dir()

    ingest_documents(doc, COLLECTION_NAME)

if __name__ == "__main__":
    build_index()
